from __future__ import annotations

import json
import base64
import os
import mimetypes
import re
import zipfile
from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Literal

from fastapi import FastAPI, HTTPException
from fastapi.encoders import jsonable_encoder
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field

from bwa_backend import app as graph_app
from bwa_backend.mcp_client import MCPError, invoke_mcp_tool, list_mcp_tools
from bwa_backend.reducer import save_blog_manifest
from bwa_backend.worker import worker_node
from bwa_backend.storage import get_blog, import_blog_from_folder, init_db, list_blogs, save_blog_markdown


ROOT_DIR = Path(__file__).resolve().parent
OUTPUTS_DIR = ROOT_DIR / "outputs"

api = FastAPI(title="Blog Writing Agent API")

api.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_origin_regex=r"^https?://(localhost|127\.0\.0\.1|\[::1\])(:\d+)?$",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

OUTPUTS_DIR.mkdir(exist_ok=True)
api.mount("/outputs", StaticFiles(directory=OUTPUTS_DIR), name="outputs")


def _bootstrap_storage() -> None:
    init_db()
    for path in OUTPUTS_DIR.glob("*/blog.md"):
        slug = path.parent.name
        try:
            existing = get_blog(slug)
            if existing:
                continue
            md_text = path.read_text(encoding="utf-8", errors="replace")
            manifest_path = path.parent / ".blog.json"
            manifest = {}
            if manifest_path.exists():
                try:
                    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
                except Exception:
                    manifest = {}
            import_blog_from_folder(slug, md_text, manifest)
        except Exception:
            continue


_bootstrap_storage()


class GenerateRequest(BaseModel):
    topic: str = Field(..., min_length=1)
    as_of: Optional[str] = None
    audience: str = "technical readers"
    tone: str = "clear and practical"
    blog_length: Literal["short", "medium", "long"] = "medium"
    generate_images: bool = True
    max_images: int = Field(3, ge=0, le=3)


class BlogSummary(BaseModel):
    filename: str
    title: str
    modified_at: float
    output_slug: Optional[str] = None


class RegenerateSectionRequest(BaseModel):
    output_slug: str = Field(..., min_length=1)
    section_id: int = Field(..., ge=1)
    topic: Optional[str] = None
    as_of: Optional[str] = None
    mode: Optional[str] = None
    recency_days: Optional[int] = None
    plan: Optional[Dict[str, Any]] = None
    evidence: List[Dict[str, Any]] = Field(default_factory=list)
    sections: List[List[Any]] = Field(default_factory=list)
    final: Optional[str] = None


class ExportRequest(BaseModel):
    markdown: str = Field(..., min_length=1)
    filename: str = Field(default="blog.md")
    output_slug: Optional[str] = None
    title: Optional[str] = None


class SaveBlogRequest(BaseModel):
    markdown: str = Field(..., min_length=1)
    title: Optional[str] = None


class ShareRequest(BaseModel):
    markdown: str = Field(..., min_length=1)
    title: Optional[str] = None
    output_slug: Optional[str] = None


def safe_slug(title: str) -> str:
    s = title.strip().lower()
    s = re.sub(r"[^a-z0-9 _-]+", "", s)
    s = re.sub(r"\s+", "_", s).strip("_")
    return s or "blog"


def extract_title_from_md(md: str, fallback: str) -> str:
    for line in md.splitlines():
        if line.startswith("# "):
            title = line[2:].strip()
            return title or fallback
    return fallback


def list_past_blogs() -> List[Path]:
    files = [p for p in OUTPUTS_DIR.glob("*/blog.md") if p.is_file()]
    files.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    return files


def _initial_graph_state(
    topic: str,
    as_of: str,
    audience: str = "technical readers",
    tone: str = "clear and practical",
    blog_length: str = "medium",
) -> Dict[str, Any]:
    return {
        "topic": topic.strip(),
        "audience": audience,
        "tone": tone,
        "blog_length": blog_length,
        "mode": "",
        "needs_research": False,
        "queries": [],
        "evidence": [],
        "plan": None,
        "as_of": as_of,
        "recency_days": 7,
        "sections": [],
        "generate_images": True,
        "max_images": 3,
        "merged_md": "",
        "md_with_placeholders": "",
        "image_specs": [],
        "final": "",
        "output_slug": "",
        "output_dir": "",
    }


def _markdown_filename(result: Dict[str, Any]) -> str:
    final_md = result.get("final") or ""
    plan = result.get("plan")
    if hasattr(plan, "blog_title"):
        title = plan.blog_title
    elif isinstance(plan, dict):
        title = plan.get("blog_title") or extract_title_from_md(final_md, "blog")
    else:
        title = extract_title_from_md(final_md, "blog")
    return "blog.md"


def _rebuild_markdown(blog_title: str, sections: List[tuple[int, str]]) -> str:
    ordered_sections = [md for _, md in sorted(sections, key=lambda x: x[0])]
    body = "\n\n".join(ordered_sections).strip()
    return f"# {blog_title}\n\n{body}\n"


def _replace_section_block(markdown: str, section_title: str, replacement: str) -> Optional[str]:
    heading = f"## {section_title}".strip()
    lines = markdown.splitlines()
    start = None
    for idx, line in enumerate(lines):
        if line.strip() == heading:
            start = idx
            break
    if start is None:
        return None

    end = len(lines)
    for idx in range(start + 1, len(lines)):
        if lines[idx].startswith("## "):
            end = idx
            break

    new_lines = lines[:start] + replacement.rstrip().splitlines() + lines[end:]
    return "\n".join(new_lines).rstrip() + "\n"


def _bundle_zip(md_text: str, md_filename: str, output_slug: Optional[str] = None) -> bytes:
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as z:
        z.writestr(md_filename, md_text.encode("utf-8"))
        if output_slug:
            images_dir = OUTPUTS_DIR / output_slug / "images"
            base_dir = OUTPUTS_DIR / output_slug
            for p in images_dir.rglob("*") if images_dir.exists() else []:
                if p.is_file():
                    z.write(p, arcname=str(p.relative_to(base_dir)))
    return buf.getvalue()


def _image_bytes_for_export(src: str, output_slug: Optional[str]) -> tuple[Optional[bytes], Optional[str]]:
    src = (src or "").strip()
    if not src or src.startswith("http://") or src.startswith("https://"):
        return None, None

    candidates = []
    if output_slug:
        candidates.append((OUTPUTS_DIR / output_slug / src.replace("./", "").lstrip("/")).resolve())
        candidates.append((OUTPUTS_DIR / output_slug / "images" / Path(src).name).resolve())
    candidates.append((ROOT_DIR / src.replace("./", "").lstrip("/")).resolve())

    for path in candidates:
        try:
            if path.exists() and path.is_file():
                return path.read_bytes(), path.suffix.lower()
        except Exception:
            continue
    return None, None


def _image_data_uri(src: str, output_slug: Optional[str]) -> Optional[str]:
    data, ext = _image_bytes_for_export(src, output_slug)
    if not data:
        return None
    mime = mimetypes.types_map.get(ext or "", "image/png")
    return f"data:{mime};base64,{base64.b64encode(data).decode('ascii')}"


def _markdown_title(md_text: str, fallback: str = "blog") -> str:
    return extract_title_from_md(md_text, fallback)


def _export_html(md_text: str, output_slug: Optional[str], title: str) -> str:
    def inline_image(match: re.Match[str]) -> str:
        alt = match.group("alt") or ""
        src = match.group("src") or ""
        data_uri = _image_data_uri(src, output_slug)
        if data_uri:
          return f'<img src="{data_uri}" alt="{alt}" />'
        return f'<img src="{src}" alt="{alt}" />'

    html = md_text
    html = re.sub(r"!\[(?P<alt>[^\]]*)\]\((?P<src>[^)]+)\)", inline_image, html)
    html = html.replace("&", "&amp;").replace("<img ", "__IMG_OPEN__").replace(" />", " />__IMG_CLOSE__")
    html = html.replace("<", "&lt;").replace(">", "&gt;")
    html = html.replace("__IMG_OPEN__", "<img ").replace("__IMG_CLOSE__", " />")
    html = html.replace("\r\n", "\n")

    lines = html.split("\n")
    out: list[str] = []
    in_ul = False
    in_ol = False
    in_code = False
    code_lines: list[str] = []

    def flush_paragraph(buffer: list[str]) -> None:
        if buffer:
            out.append(f"<p>{' '.join(buffer).strip()}</p>")
            buffer.clear()

    paragraph: list[str] = []
    for raw_line in lines:
        line = raw_line.rstrip()
        if line.startswith("```"):
            if in_code:
                out.append("<pre><code>" + "\n".join(code_lines) + "</code></pre>")
                code_lines = []
                in_code = False
            else:
                flush_paragraph(paragraph)
                if in_ul:
                    out.append("</ul>")
                    in_ul = False
                if in_ol:
                    out.append("</ol>")
                    in_ol = False
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            flush_paragraph(paragraph)
            if in_ul:
                out.append("</ul>")
                in_ul = False
            if in_ol:
                out.append("</ol>")
                in_ol = False
            continue

        if line.startswith("# "):
            flush_paragraph(paragraph)
            out.append(f"<h1>{line[2:].strip()}</h1>")
            continue
        if line.startswith("## "):
            flush_paragraph(paragraph)
            out.append(f"<h2>{line[3:].strip()}</h2>")
            continue
        if line.startswith("### "):
            flush_paragraph(paragraph)
            out.append(f"<h3>{line[4:].strip()}</h3>")
            continue
        if line.startswith("- "):
            flush_paragraph(paragraph)
            if in_ol:
                out.append("</ol>")
                in_ol = False
            if not in_ul:
                out.append("<ul>")
                in_ul = True
            out.append(f"<li>{line[2:].strip()}</li>")
            continue
        if re.match(r"^\d+\.\s+", line):
            flush_paragraph(paragraph)
            if in_ul:
                out.append("</ul>")
                in_ul = False
        if not in_ol:
            out.append("<ol>")
            in_ol = True
            numbered_text = re.sub(r"^\d+\.\s+", "", line).strip()
            out.append(f"<li>{numbered_text}</li>")
            continue
        if line.startswith("> "):
            flush_paragraph(paragraph)
            out.append(f"<blockquote>{line[2:].strip()}</blockquote>")
            continue
        paragraph.append(line.strip())

    flush_paragraph(paragraph)
    if in_ul:
        out.append("</ul>")
    if in_ol:
        out.append("</ol>")
    if in_code:
        out.append("<pre><code>" + "\n".join(code_lines) + "</code></pre>")

    body = "\n".join(out)
    return f"""<!doctype html>
<html>
<head>
  <meta charset="utf-8" />
  <title>{title}</title>
  <style>
    body {{ font-family: system-ui, -apple-system, Segoe UI, sans-serif; line-height: 1.6; max-width: 900px; margin: 40px auto; padding: 0 20px; color: #1d252f; }}
    h1, h2, h3 {{ line-height: 1.2; }}
    img {{ max-width: 100%; height: auto; display: block; margin: 18px 0; }}
    pre {{ background: #18201d; color: #f4faf6; padding: 14px; border-radius: 8px; overflow: auto; }}
    code {{ font-family: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, monospace; }}
    blockquote {{ border-left: 4px solid #dfe3dd; margin: 12px 0; padding: 6px 14px; color: #4a5550; }}
  </style>
</head>
<body>
{body}
</body>
</html>"""


def _export_docx(md_text: str, output_slug: Optional[str], title: str) -> bytes:
    from docx import Document
    from docx.shared import Inches

    document = Document()
    document.add_heading(title, level=0)

    lines = md_text.splitlines()
    paragraph: list[str] = []
    in_code = False
    code_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph
        if paragraph:
            document.add_paragraph(" ".join(paragraph).strip())
            paragraph = []

    for line in lines:
        line = line.rstrip()
        if line.startswith("```"):
            if in_code:
                flush_paragraph()
                document.add_paragraph("\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                flush_paragraph()
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            flush_paragraph()
            continue
        if line.startswith("# "):
            flush_paragraph()
            document.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            flush_paragraph()
            document.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            flush_paragraph()
            document.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("- "):
            flush_paragraph()
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        if re.match(r"^\d+\.\s+", line):
            flush_paragraph()
            document.add_paragraph(re.sub(r"^\d+\.\s+", "", line).strip(), style="List Number")
            continue
        img_match = re.match(r"!\[(?P<alt>[^\]]*)\]\((?P<src>[^)]+)\)", line.strip())
        if img_match:
            flush_paragraph()
            data, _ext = _image_bytes_for_export(img_match.group("src"), output_slug)
            if data:
                tmp = BytesIO(data)
                try:
                    document.add_picture(tmp, width=Inches(5.8))
                except Exception:
                    document.add_paragraph(img_match.group("alt") or "Image")
            else:
                document.add_paragraph(img_match.group("alt") or "Image")
            continue
        paragraph.append(line.strip())

    flush_paragraph()
    if in_code and code_lines:
        document.add_paragraph("\n".join(code_lines))

    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()


def _export_pdf(md_text: str, output_slug: Optional[str], title: str) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Image, Paragraph, Spacer, SimpleDocTemplate

    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles["Title"]), Spacer(1, 0.2 * inch)]
    lines = md_text.splitlines()
    paragraph: list[str] = []
    in_code = False
    code_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph
        if paragraph:
            story.append(Paragraph(" ".join(paragraph).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), styles["BodyText"]))
            story.append(Spacer(1, 0.12 * inch))
            paragraph = []

    for line in lines:
        line = line.rstrip()
        if line.startswith("```"):
            if in_code:
                flush_paragraph()
                story.append(Paragraph("<font name=\"Courier\">" + "<br/>".join([ln.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;") for ln in code_lines]) + "</font>", styles["Code"]))
                story.append(Spacer(1, 0.12 * inch))
                code_lines = []
                in_code = False
            else:
                flush_paragraph()
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            flush_paragraph()
            continue
        if line.startswith("# "):
            flush_paragraph()
            story.append(Paragraph(line[2:].strip(), styles["Heading1"]))
            story.append(Spacer(1, 0.12 * inch))
            continue
        if line.startswith("## "):
            flush_paragraph()
            story.append(Paragraph(line[3:].strip(), styles["Heading2"]))
            story.append(Spacer(1, 0.1 * inch))
            continue
        if line.startswith("### "):
            flush_paragraph()
            story.append(Paragraph(line[4:].strip(), styles["Heading3"]))
            story.append(Spacer(1, 0.08 * inch))
            continue
        if line.startswith("- "):
            flush_paragraph()
            story.append(Paragraph(f"• {line[2:].strip()}", styles["BodyText"]))
            continue
        if re.match(r"^\d+\.\s+", line):
            flush_paragraph()
            story.append(Paragraph(re.sub(r"^\d+\.\s+", "", line).strip(), styles["BodyText"]))
            continue
        img_match = re.match(r"!\[(?P<alt>[^\]]*)\]\((?P<src>[^)]+)\)", line.strip())
        if img_match:
            flush_paragraph()
            data, _ext = _image_bytes_for_export(img_match.group("src"), output_slug)
            if data:
                try:
                    story.append(Image(BytesIO(data), width=5.8 * inch, preserveAspectRatio=True))
                    story.append(Spacer(1, 0.12 * inch))
                except Exception:
                    story.append(Paragraph(img_match.group("alt") or "Image", styles["BodyText"]))
            else:
                story.append(Paragraph(img_match.group("alt") or "Image", styles["BodyText"]))
            continue
        paragraph.append(line.strip())

    flush_paragraph()
    if in_code and code_lines:
        story.append(Paragraph("<font name=\"Courier\">" + "<br/>".join(code_lines) + "</font>", styles["Code"]))

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, title=title)
    doc.build(story)
    return buf.getvalue()


def _export_payload(md_text: str, filename: str, output_slug: Optional[str], title: Optional[str], kind: str) -> tuple[bytes, str, str]:
    safe_base = safe_slug(Path(filename).stem or "blog")
    ext_map = {"html": "html", "docx": "docx", "pdf": "pdf"}
    out_title = title or _markdown_title(md_text, safe_base)
    if kind == "html":
        data = _export_html(md_text, output_slug, out_title).encode("utf-8")
        return data, "text/html", f"{safe_base}.html"
    if kind == "docx":
        return _export_docx(md_text, output_slug, out_title), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{safe_base}.docx"
    if kind == "pdf":
        return _export_pdf(md_text, output_slug, out_title), "application/pdf", f"{safe_base}.pdf"
    raise HTTPException(status_code=400, detail="Unsupported export format")


def _pick_tool_name(available: List[dict], preferred: Optional[str], candidates: List[str]) -> str:
    names = [str(item.get("name")) for item in available if item.get("name")]
    if preferred:
        if preferred in names or not names:
            return preferred
    for candidate in candidates:
        if candidate in names:
            return candidate
    if names:
        return names[0]
    raise HTTPException(status_code=404, detail="No MCP tools available")


def _share_highlights(markdown: str, limit: int = 4) -> List[str]:
    highlights: List[str] = []
    for line in markdown.splitlines():
        item = line.strip()
        if not item:
            continue
        if item.startswith("#") or item.startswith("```") or item.startswith("> "):
            continue
        if item.startswith(("- ", "* ")):
            highlights.append(item[2:].strip())
        elif re.match(r"^\d+\.\s+", item):
            highlights.append(re.sub(r"^\d+\.\s+", "", item).strip())
        if len(highlights) >= limit:
            break
    return [item for item in highlights if item]


def _share_excerpt(markdown: str, limit: int = 700) -> str:
    parts: List[str] = []
    for line in markdown.splitlines():
        item = line.strip()
        if not item or item.startswith("#") or item.startswith("```"):
            continue
        if item.startswith("> "):
            continue
        parts.append(item)
        if len(" ".join(parts)) >= limit:
            break
    excerpt = " ".join(parts).strip()
    if len(excerpt) > limit:
        excerpt = excerpt[:limit].rstrip() + "..."
    return excerpt


def _share_slack_blocks(title: str, output_slug: str, markdown: str) -> List[Dict[str, Any]]:
    highlights = _share_highlights(markdown, limit=3)
    excerpt = _share_excerpt(markdown, limit=450)
    block_text = f"*{title}*\nA fresh blog draft is ready for review."
    fields = [
        {"type": "mrkdwn", "text": f"*Status*\nReady to review"},
        {"type": "mrkdwn", "text": f"*Slug*\n`{output_slug}`"},
        {"type": "mrkdwn", "text": f"*Highlights*\n{len(highlights)} sections"},
    ]
    blocks: List[Dict[str, Any]] = [
        {"type": "header", "text": {"type": "plain_text", "text": title[:150]}},
        {"type": "section", "text": {"type": "mrkdwn", "text": block_text}},
        {"type": "section", "fields": fields},
        {"type": "divider"},
    ]
    if highlights:
        blocks.append(
            {
                "type": "section",
                "text": {"type": "mrkdwn", "text": "*Top takeaways*\n" + "\n".join(f"• {item}" for item in highlights)},
            }
        )
    if excerpt:
        blocks.append({"type": "section", "text": {"type": "mrkdwn", "text": f"*Preview*\n{excerpt}"}})
    blocks.append({"type": "context", "elements": [{"type": "mrkdwn", "text": "Use the blog app to edit, export, or regenerate sections."}]})
    return blocks


def _share_gmail_html(title: str, output_slug: str, markdown: str) -> str:
    intro = _share_excerpt(markdown, limit=360)
    rendered = _export_html(markdown, output_slug, title)
    body_match = re.search(r"<body>(.*)</body>", rendered, flags=re.S)
    body_content = body_match.group(1).strip() if body_match else rendered
    return f"""<!doctype html>
<html>
<head>
  <meta charset="utf-8" />
  <style>
    body {{ margin: 0; background: #f6f8fb; color: #1d252f; font-family: system-ui, -apple-system, Segoe UI, sans-serif; }}
    .wrap {{ max-width: 900px; margin: 0 auto; padding: 28px 20px 40px; }}
    .hero {{ background: linear-gradient(135deg, #1f3a5f, #2e4f7d); color: #fff; border-radius: 16px; padding: 24px 26px; }}
    .hero h1 {{ margin: 0 0 10px; font-size: 30px; line-height: 1.15; }}
    .meta {{ margin: 0; opacity: 0.9; font-size: 14px; }}
    .intro {{ margin: 16px 0 0; font-size: 16px; line-height: 1.7; }}
    .card {{ background: #fff; border-radius: 16px; padding: 22px 26px; margin-top: 18px; box-shadow: 0 10px 30px rgba(18, 32, 53, 0.08); }}
    .badge {{ display: inline-block; background: #e8eef7; color: #244064; border-radius: 999px; padding: 6px 12px; font-size: 12px; font-weight: 700; margin-top: 10px; }}
    .footer {{ margin-top: 20px; color: #637085; font-size: 13px; }}
  </style>
</head>
<body>
  <div class="wrap">
    <div class="hero">
      <div class="badge">Blog draft ready</div>
      <h1>{title}</h1>
      <p class="meta">Slug: {output_slug}</p>
      <p class="intro">{intro}</p>
    </div>
    <div class="card">
      {body_content}
    </div>
    <p class="footer">Open the blog app to edit sections, export, or send another version.</p>
  </div>
</body>
</html>"""


def _share_via_mcp(destination: str, payload: ShareRequest) -> dict:
    title = payload.title or extract_title_from_md(payload.markdown, "blog")
    output_slug = payload.output_slug or safe_slug(title)

    try:
        with_path = {}
        if destination == "gmail":
            html_body = _share_gmail_html(title, output_slug, payload.markdown)
            with_path = {
                "subject": f"Draft ready: {title}",
                "body": _share_excerpt(payload.markdown, limit=1200),
                "html_body": html_body,
                "to": (os.getenv("GMAIL_TO") or "").strip() or None,
                "output_slug": output_slug,
            }
            preferred_tool = (os.getenv("GMAIL_MCP_TOOL") or "").strip() or None
            candidates = ["send_email", "draft_email", "gmail_send", "compose_email"]
        elif destination == "slack":
            blocks = _share_slack_blocks(title, output_slug, payload.markdown)
            excerpt = _share_excerpt(payload.markdown, limit=500)
            with_path = {
                "channel": (os.getenv("SLACK_CHANNEL") or "").strip() or None,
                "text": excerpt,
                "title": f"Blog draft ready: {title}",
                "output_slug": output_slug,
                "blocks": blocks,
            }
            preferred_tool = (os.getenv("SLACK_MCP_TOOL") or "").strip() or None
            candidates = ["post_message", "send_message", "chat_postMessage", "slack_post_message"]
        else:
            raise HTTPException(status_code=400, detail="Unsupported destination")

        tool = _pick_tool_name(list_mcp_tools(destination), preferred_tool, candidates)
        result = invoke_mcp_tool(destination, tool, with_path)
        return {
            "ok": True,
            "destination": destination,
            "tool": tool,
            "output_slug": output_slug,
            "result": result.raw,
        }
    except MCPError as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from exc


def _event(event: str, data: Dict[str, Any]) -> str:
    return f"event: {event}\ndata: {json.dumps(jsonable_encoder(data))}\n\n"


def _summary(state: Dict[str, Any], status: str, message: str) -> Dict[str, Any]:
    plan = state.get("plan")
    tasks = []
    if hasattr(plan, "tasks"):
        tasks = plan.tasks
    elif isinstance(plan, dict):
        tasks = plan.get("tasks") or []

    return {
        "status": status,
        "message": message,
        "mode": state.get("mode"),
        "needs_research": state.get("needs_research"),
        "queries_count": len(state.get("queries") or []),
        "evidence_count": len(state.get("evidence") or []),
        "tasks_count": len(tasks),
        "sections_done": len(state.get("sections") or []),
        "images_count": len(state.get("image_specs") or []),
        "has_final": bool(state.get("final")),
    }


def _progress_message(state: Dict[str, Any]) -> str:
    if state.get("final"):
        return "Blog complete"
    if state.get("image_specs"):
        return "Images planned"
    if state.get("md_with_placeholders"):
        return "Image placeholders placed"
    if state.get("merged_md"):
        return "Sections merged"
    if state.get("sections"):
        return f"Sections generated: {len(state.get('sections') or [])}"
    if state.get("plan"):
        return "Plan ready"
    if state.get("needs_research") and state.get("evidence"):
        return "Research done"
    if state.get("needs_research") is False and state.get("mode"):
        return "Router done; research skipped"
    if state.get("mode"):
        return "Router done"
    return "Starting graph"


def _stream_graph(inputs: Dict[str, Any]):
    latest: Dict[str, Any] = {}
    yield _event("progress", _summary(latest, "running", "Starting graph"))

    try:
        for state in graph_app.stream(inputs, stream_mode="values"):
            latest = state
            message = _progress_message(latest)
            yield _event(
                "progress",
                {
                    **_summary(latest, "running", message),
                    "state": latest,
                },
            )

        payload = jsonable_encoder(latest)
        payload["markdown_filename"] = _markdown_filename(latest)
        yield _event(
            "complete",
            {
                **_summary(latest, "complete", "Generation complete"),
                "result": payload,
            },
        )
    except Exception as exc:
        yield _event("error", {"status": "error", "message": str(exc)})


@api.post("/api/blogs/{output_slug}/sections/{section_id}/regenerate")
def regenerate_section(output_slug: str, section_id: int, request: RegenerateSectionRequest) -> dict:
    if request.output_slug != output_slug:
        raise HTTPException(status_code=400, detail="output_slug mismatch")

    record = get_blog(output_slug)
    if not record:
        raise HTTPException(status_code=404, detail="Blog not found")

    plan = request.plan or record.get("plan")
    if not plan:
        raise HTTPException(status_code=404, detail="Blog metadata not found")

    tasks = plan.get("tasks") or []
    task = next((item for item in tasks if int(item.get("id", 0)) == section_id), None)
    if not task:
        raise HTTPException(status_code=404, detail="Section task not found")

    sections_src = request.sections or record.get("sections") or []
    sections = [(int(item[0]), str(item[1])) for item in sections_src if len(item) >= 2]
    previous_md = next((md for sid, md in sections if sid == section_id), None)
    if previous_md is None:
        raise HTTPException(status_code=404, detail="Current section markdown not found")

    topic = request.topic or record.get("topic") or plan.get("blog_title") or output_slug
    mode = request.mode or record.get("mode") or "closed_book"
    recency_days = int(request.recency_days or record.get("recency_days") or 3650)
    evidence = request.evidence or record.get("evidence") or []
    as_of = request.as_of or record.get("as_of") or date.today().isoformat()

    regenerated = worker_node(
        {
            "task": task,
            "topic": topic,
            "mode": mode,
            "as_of": as_of,
            "recency_days": recency_days,
            "plan": plan,
            "evidence": evidence,
            "generate_images": record.get("generate_images", True),
            "max_images": record.get("max_images", 3),
        }
    )
    new_md = str(regenerated.get("sections", [[section_id, previous_md]])[0][1])

    updated_sections: List[tuple[int, str]] = []
    replaced = False
    for sid, md in sections:
        if sid == section_id and not replaced:
            updated_sections.append((sid, new_md))
            replaced = True
        else:
            updated_sections.append((sid, md))

    blog_title = str(plan.get("blog_title") or output_slug)
    current_final = str(request.final or record.get("final") or "")
    if current_final:
        final_md = _replace_section_block(current_final, str(task.get("title") or section_id), new_md) or current_final
    else:
        final_md = _rebuild_markdown(blog_title, updated_sections)
    save_blog_manifest(
        blog_title,
        final_md,
        {
            **record,
            "plan": plan,
            "evidence": evidence,
            "sections": updated_sections,
            "generate_images": record.get("generate_images", True),
            "max_images": record.get("max_images", 3),
        },
    )

    payload = {
        "output_slug": output_slug,
        "output_dir": f"outputs/{output_slug}",
        "final": final_md,
        "sections": updated_sections,
        "markdown_filename": "blog.md",
    }
    return payload


@api.get("/api/health")
def health() -> dict:
    return {"ok": True}


@api.post("/api/generate")
def generate_blog(request: GenerateRequest) -> dict:
    as_of = request.as_of or date.today().isoformat()
    inputs = _initial_graph_state(request.topic, as_of, request.audience, request.tone, request.blog_length)
    inputs["generate_images"] = request.generate_images
    inputs["max_images"] = request.max_images
    result = graph_app.invoke(inputs)
    payload = jsonable_encoder(result)
    payload["markdown_filename"] = _markdown_filename(result)
    return payload


@api.post("/api/generate/stream")
def generate_blog_stream(request: GenerateRequest) -> StreamingResponse:
    as_of = request.as_of or date.today().isoformat()
    inputs = _initial_graph_state(request.topic, as_of, request.audience, request.tone, request.blog_length)
    inputs["generate_images"] = request.generate_images
    inputs["max_images"] = request.max_images
    return StreamingResponse(
        _stream_graph(inputs),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@api.get("/api/blogs", response_model=List[BlogSummary])
def blogs() -> List[BlogSummary]:
    return [BlogSummary(**item) for item in list_blogs()]


@api.get("/api/blogs/{filename}")
def blog(filename: str) -> dict:
    record = get_blog(filename)
    if not record:
        raise HTTPException(status_code=404, detail="Blog not found")
    return record


@api.post("/api/bundle")
def bundle(payload: Dict[str, Any]) -> Response:
    md_text = str(payload.get("markdown") or "")
    output_slug = str(payload.get("output_slug") or "").strip() or None
    raw_filename = Path(str(payload.get("filename") or "blog.md")).name
    filename = safe_slug(raw_filename.removesuffix(".md"))
    if not filename.endswith(".md"):
        filename = f"{filename}.md"
    data = _bundle_zip(md_text, filename, output_slug)
    return Response(
        content=data,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{filename.removesuffix(".md")}_bundle.zip"'},
    )


@api.post("/api/export/{kind}")
def export(kind: str, payload: ExportRequest) -> Response:
    data, media_type, download_name = _export_payload(
        payload.markdown,
        payload.filename,
        payload.output_slug,
        payload.title,
        kind.lower(),
    )
    return Response(
        content=data,
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{download_name}"'},
    )


@api.put("/api/blogs/{output_slug}")
def save_blog(output_slug: str, payload: SaveBlogRequest) -> dict:
    record = save_blog_markdown(output_slug, payload.markdown, payload.title)
    if not record:
        raise HTTPException(status_code=404, detail="Blog not found")
    return record


@api.post("/api/share/gmail")
def share_gmail(payload: ShareRequest) -> dict:
    return _share_via_mcp("gmail", payload)


@api.post("/api/share/slack")
def share_slack(payload: ShareRequest) -> dict:
    return _share_via_mcp("slack", payload)
